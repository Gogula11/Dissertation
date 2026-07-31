#!/usr/bin/env python3
"""Build dissertation_draft.docx from chapter markdown files."""
import re, os, sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# --- LaTeX to Unicode converter ---

GREEK = {
    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
    'varepsilon': 'ε', 'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι',
    'kappa': 'κ', 'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ',
    'pi': 'π', 'rho': 'ρ', 'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ',
    'phi': 'φ', 'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
    'Gamma': 'Γ', 'Delta': 'Δ', 'Theta': 'Θ', 'Lambda': 'Λ',
    'Xi': 'Ξ', 'Pi': 'Π', 'Sigma': 'Σ', 'Phi': 'Φ', 'Psi': 'Ψ', 'Omega': 'Ω',
}

SUPERSCRIPT = str.maketrans('0123456789+-=()ni', '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿⁱ')
SUBSCRIPT = str.maketrans('0123456789+-=()aejknprstux', '₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑⱼₖₙₚᵣₛₜᵤₓ')

def latex_to_unicode(s):
    """Convert a LaTeX math string to Unicode text."""
    # Remove \text{...} wrapper
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)

    # Escaped underscore: \_ → _ (must be before subscript conversion)
    s = s.replace('\\_', '_')

    # \frac{num}{den} → num/den
    while '\\frac' in s:
        m = re.search(r'\\frac\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', s)
        if not m:
            break
        num = m.group(1).strip()
        den = m.group(2).strip()
        s = s[:m.start()] + f'{num}/{den}' + s[m.end():]

    # \pmod{x} or \pmod x → (mod x) — must come before \pm replacement
    s = re.sub(r'\\pmod\{([^}]*)\}', r'(mod \1)', s)
    s = re.sub(r'\\pmod\s+(\w)', r'(mod \1)', s)
    s = re.sub(r'\\mod\b', ' mod', s)

    # Greek letters
    for name, char in GREEK.items():
        s = s.replace(f'\\{name}', char)

    # Symbols — order matters: \pmod already handled, \pm is safe
    s = s.replace('\\times', '×')
    s = s.replace('\\cdot', '·')
    s = s.replace('\\dots', '…')
    s = s.replace('\\ge', '≥')
    s = s.replace('\\le', '≤')
    s = s.replace('\\neq', '≠')
    s = s.replace('\\approx', '≈')
    s = s.replace('\\infty', '∞')
    s = s.replace('\\pm', '±')
    s = s.replace('\\rightarrow', '→')
    s = s.replace('\\leftarrow', '←')
    s = s.replace('\\Rightarrow', '⇒')
    s = s.replace('\\Leftarrow', '⇐')
    s = s.replace('\\% ', '%')
    s = s.replace('\\%', '%')

    # Superscripts: ^{...} or ^x
    def sup_repl(m):
        content = m.group(1) if m.group(1) else m.group(2)
        return content.translate(SUPERSCRIPT)
    s = re.sub(r'\^\{([^}]*)\}|\^(\w)', sup_repl, s)

    # Subscripts: _{...} only (bare _x would match literal underscores)
    def sub_repl(m):
        content = m.group(1)
        return content.translate(SUBSCRIPT)
    s = re.sub(r'_\{([^}]*)\}', sub_repl, s)

    # Clean up remaining backslashes
    s = s.replace('\\', '')

    # Clean up extra braces
    s = s.replace('{', '').replace('}', '')

    return s

# --- DOCX building ---

files = [
    ("notes/abstract.md", "Abstract"),
    ("notes/chapter_01_introduction.md", None),
    ("notes/chapter_02_background.md", None),
    ("notes/chapter_03_system_design.md", None),
    ("notes/chapter_04_implementation_results.md", None),
    ("notes/chapter_05_evaluation.md", None),
    ("notes/chapter_06_conclusions.md", None),
]

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_inline(paragraph, text):
    """Bold, inline code, and LaTeX math."""
    remaining = text
    while remaining:
        bold_m = re.search(r'\*\*(.*?)\*\*', remaining)
        code_m = re.search(r'`(.*?)`', remaining)
        math_m = re.search(r'\$(.*?)\$', remaining)
        matches = []
        if bold_m: matches.append(('bold', bold_m))
        if code_m: matches.append(('code', code_m))
        if math_m: matches.append(('math', math_m))
        if not matches:
            paragraph.add_run(remaining)
            break
        matches.sort(key=lambda x: x[1].start())
        mtype, m = matches[0]
        if m.start() > 0:
            paragraph.add_run(remaining[:m.start()])
        if mtype == 'bold':
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif mtype == 'code':
            run = paragraph.add_run(m.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif mtype == 'math':
            converted = latex_to_unicode(m.group(1))
            run = paragraph.add_run(converted)
            run.font.name = 'Cambria Math'
            run.italic = True
        remaining = remaining[m.end():]

def clean_heading(text):
    return re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()

def is_table_sep(line):
    return bool(re.match(r'^\|[\s\-:]+(\|[\s\-:]+)+\|$', line.strip()))

def parse_row(line):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'): s = s[:-1]
    return [c.strip() for c in s.split('|')]

def is_flowchart_line(stripped):
    if re.match(r'^[\s]*[\+][\s\-+]+[\+]', stripped): return True
    if re.match(r'^[\s]*[v^]\s', stripped): return True
    if re.match(r'^[\s]+v\s+v', stripped): return True
    if re.match(r'^[\s]*\^[\s]*$', stripped): return True
    if re.match(r'^[\s]*SPT\s', stripped): return True
    if re.match(r'^[\s]*\(heuristics', stripped): return True
    if re.match(r'^[\s]*\(ga\.py\)', stripped): return True
    if re.match(r'^[\s]*\(ga_env', stripped): return True
    if re.match(r'^[\s]*\|$', stripped): return True
    return False

figure_list = []

for filepath, title_override in files:
    full_path = os.path.join(ROOT, filepath)
    with open(full_path) as f:
        text = f.read()
    lines = text.split('\n')
    i = 0
    in_code_block = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1; continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            i += 1; continue
        if not stripped:
            i += 1; continue
        if is_flowchart_line(stripped):
            i += 1; continue
        if stripped.startswith('#### '):
            doc.add_heading(clean_heading(stripped[5:]), level=3)
            i += 1; continue
        if stripped.startswith('### '):
            doc.add_heading(clean_heading(stripped[4:]), level=2)
            i += 1; continue
        if stripped.startswith('## '):
            doc.add_heading(clean_heading(stripped[3:]), level=1)
            i += 1; continue
        if stripped.startswith('# '):
            doc.add_heading(clean_heading(stripped[2:]), level=0)
            i += 1; continue
        m = re.match(r'^\*\*Chapter (\d+)\.\s*(.+?)\*\*$', stripped)
        if m:
            doc.add_heading(f'Chapter {m.group(1)}. {m.group(2)}', level=0)
            i += 1; continue
        m = re.match(r'^\*\*(\d+(?:\.\d+)*)\s+(.+?)\*\*$', stripped)
        if m:
            depth = m.group(1).count('.')
            level = min(depth + 1, 3)
            doc.add_heading(f'{m.group(1)} {m.group(2)}', level=level)
            i += 1; continue
        img_m = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if img_m:
            caption = img_m.group(1)
            img_path = img_m.group(2)
            full_img = None
            for base in [os.path.join(ROOT, 'notes'), ROOT]:
                candidate = os.path.normpath(os.path.join(base, img_path))
                if os.path.exists(candidate):
                    full_img = candidate
                    break
            if full_img:
                try:
                    doc.add_picture(full_img, width=Inches(5.5))
                except:
                    doc.add_paragraph(f"[Image error: {caption}]")
            else:
                doc.add_paragraph(f"[Missing: {img_path}]")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(9)
            run.italic = True
            figure_list.append(caption)
            i += 1; continue
        if stripped.startswith('|') and not is_table_sep(stripped):
            table_rows = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith('|'):
                    if is_table_sep(s):
                        i += 1; continue
                    table_rows.append(parse_row(s))
                    i += 1
                else:
                    break
            if table_rows:
                n_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=n_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < n_cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_inline(p, cell_text)
                            for run in p.runs:
                                run.font.size = Pt(9)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph()
            continue
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:])
            i += 1; continue
        num_m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_m:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, num_m.group(2))
            i += 1; continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

outpath = os.path.join(ROOT, "notes", "dissertation_draft.docx")
doc.save(outpath)
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
print(f"Saved: {outpath}")
print(f"Figures: {len(figure_list)}, Embedded: {img_count}, Tables: {len(doc.tables)}")
